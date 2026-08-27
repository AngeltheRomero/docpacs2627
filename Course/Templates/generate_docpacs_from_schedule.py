from __future__ import annotations

import argparse
import calendar
import re
import shutil
import sys
from collections import OrderedDict
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from openpyxl import load_workbook
    from openpyxl.utils.datetime import from_excel
except ImportError as exc:
    missing = getattr(exc, "name", "a required package")
    raise SystemExit(
        f"Missing {missing}. Install the required packages with:\n"
        "    py -m pip install python-docx openpyxl"
    ) from exc


CODE_RE = re.compile(r"^\d{6}$")
NODEJS_RE = re.compile(
    r"^Unit\s*(\d+)\s*,\s*Lesson\s*(\d+)\s*,\s*(.+?)\s*$",
    re.IGNORECASE,
)
PROF_RE = re.compile(r"^Prof\s*:\s*(.+?)\s*$", re.IGNORECASE)


@dataclass
class Assignment:
    raw_name: str
    display_name: str
    rubric_kind: str
    sources: set[str] = field(default_factory=set)

    @property
    def who(self) -> str:
        if "certification" in self.sources:
            return "Juniors"
        if "juniors" in self.sources and "seniors" in self.sources:
            return "Everyone"
        if "juniors" in self.sources:
            return "Juniors"
        if "seniors" in self.sources:
            return "Seniors"
        return "Everyone"


def first_existing(directory: Path, names: Iterable[str]) -> Path:
    candidates = [directory / name for name in names]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def parse_args() -> argparse.Namespace:
    script_dir = Path(__file__).resolve().parent

    schedule_dir = script_dir / "Course" / "Information"
    docpac_dir = script_dir / "Course" / "Templates"

    parser = argparse.ArgumentParser(
        description=(
            "Rebuild future DocPacs from a template and schedule, populate the "
            "assignment list, and include only the rubrics required for that week."
        )
    )
    parser.add_argument(
        "--root",
        type=Path,
        default=script_dir,
        help="Folder containing docpac_XXXXXX subfolders.",
    )

    parser.add_argument(
        "--schedule",
        type=Path,
        default=first_existing(
            schedule_dir,
            ("Schedule.xlsx",),
        ),
        help="Excel schedule file.",
    )

    parser.add_argument(
        "--template",
        type=Path,
        default=first_existing(
            docpac_dir,
            ("docpac_template.docx",),
        ),
        help="Word DocPac template.",
    )
    parser.add_argument(
        "--sheet",
        default=None,
        help="Worksheet name. The active worksheet is used by default.",
    )
    parser.add_argument(
        "--as-of",
        type=date.fromisoformat,
        default=date.today(),
        metavar="YYYY-MM-DD",
        help=(
            "Date used to decide whether a DocPac is in the past. "
            "Defaults to today's date."
        ),
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show what would be rebuilt without changing files.",
    )
    return parser.parse_args()


def normalize_code(value: Any) -> str:
    if value is None:
        raise ValueError("blank item code")
    if isinstance(value, bool):
        raise ValueError(f"invalid item code: {value!r}")

    if isinstance(value, (int, float)):
        if isinstance(value, float) and not value.is_integer():
            raise ValueError(f"item code is not a whole number: {value!r}")
        code = f"{int(value):06d}"
    else:
        text = str(value).strip()
        if text.endswith(".0") and text[:-2].isdigit():
            text = text[:-2]
        code = text.zfill(6)

    if not CODE_RE.fullmatch(code):
        raise ValueError(f"item code must contain six digits: {value!r}")
    return code


def academic_date_from_code(code: str) -> date:
    day_number = int(code[2:4])
    month_number = int(code[4:6])

    if 8 <= month_number <= 12:
        year = 2026
    elif 1 <= month_number <= 5:
        year = 2027
    else:
        raise ValueError(
            f"item code {code} uses month {month_number}; expected August-May"
        )

    return date(year, month_number, day_number)


def plain_number(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def format_due_date(value: Any, workbook_epoch: datetime) -> str:
    if value is None or value == "":
        return ""

    converted: date | datetime | None = None
    if isinstance(value, (date, datetime)):
        converted = value
    elif isinstance(value, (int, float)):
        try:
            converted = from_excel(value, epoch=workbook_epoch)
        except (TypeError, ValueError, OverflowError):
            converted = None

    if converted is not None:
        return f"{converted.day}-{converted.strftime('%b')}"
    return str(value).strip()


def clean_item(text: str) -> str:
    cleaned = text.strip().strip("•").strip()
    if len(cleaned) >= 2 and cleaned[0] == cleaned[-1] and cleaned[0] in {'"', "'"}:
        cleaned = cleaned[1:-1].strip()
    return re.sub(r"\s+", " ", cleaned)


def split_assignment_cell(value: Any) -> list[str]:
    """
    Split assignments only on Excel line breaks.

    Commas inside assignment names are preserved.
    """
    if value is None:
        return []

    text = str(value).replace("\r\n", "\n").replace("\r", "\n")

    return [
        item
        for raw_line in text.split("\n")
        if (item := clean_item(raw_line))
    ]


def safe_folder_name(text: str) -> str:
    """
    Convert spreadsheet text into a valid Windows folder name.

    Invalid filename characters are replaced with underscores, and trailing
    spaces or periods are removed.
    """
    cleaned = re.sub(r'[<>:"/\\|?*]', "_", text).strip().rstrip(". ")
    if cleaned in {"", ".", ".."}:
        raise ValueError(f"invalid Coursework folder name: {text!r}")
    return cleaned


def create_coursework_folders(parent: Path, coursework_items: list[str]) -> list[Path]:
    created: list[Path] = []

    for item in coursework_items:
        folder = parent / safe_folder_name(item)
        folder.mkdir(parents=False, exist_ok=True)
        created.append(folder)
        txt_file = folder / f"{folder.name}.txt"
        txt_file.touch(exist_ok=True)

    return created


def normalized_key(text: str) -> str:
    key = clean_item(text).casefold()
    key = re.sub(r"\s+", " ", key)
    key = re.sub(r"^company projects?$", "company projects", key)
    key = re.sub(r"^performance reviews?$", "performance review", key)
    return key


def is_company_project(text: str) -> bool:
    return normalized_key(text) == "company projects"


def is_performance_review(text: str) -> bool:
    return normalized_key(text) == "performance review"


def classify_junior_item(text: str) -> tuple[str, str]:
    node_match = NODEJS_RE.fullmatch(text)
    if node_match:
        unit, lesson, name = node_match.groups()
        display = f"NodeJS Backend: Unit {int(unit)}, Lesson {int(lesson)}, {name.strip()}"
        return display, "nodejs"

    prof_match = PROF_RE.fullmatch(text)
    if prof_match:
        return (
            f"Collaboration and Professionalism: {prof_match.group(1).strip()}",
            "professional",
        )

    if is_performance_review(text):
        return "Performance Review", "performance"

    if is_company_project(text):
        return "Company Projects", "company"

    return f"Web Browser Frontend: {text}", "web"


def classify_senior_item(text: str) -> tuple[str, str]:
    if is_company_project(text):
        return "Company Projects", "company"
    return f"Senior Assignment: {text}", "senior"


def classify_certification_item(text: str) -> tuple[str, str]:
    return f"Certification: {text}", "certification"


def build_assignments(
    junior_items: list[str],
    senior_items: list[str],
    certification_items: list[str],
) -> list[Assignment]:
    merged: OrderedDict[str, Assignment] = OrderedDict()

    for source, items in (("juniors", junior_items), ("seniors", senior_items)):
        for item in items:
            key = normalized_key(item)
            if key in merged:
                merged[key].sources.add(source)
                continue

            if source == "juniors":
                display, rubric_kind = classify_junior_item(item)
            else:
                display, rubric_kind = classify_senior_item(item)

            merged[key] = Assignment(
                raw_name=item,
                display_name=display,
                rubric_kind=rubric_kind,
                sources={source},
            )

    # Certification/Coursework entries are a separate assignment category even
    # if their text happens to match something in Juniors or Seniors.
    for item in certification_items:
        key = "certification::" + normalized_key(item)
        if key in merged:
            continue
        display, rubric_kind = classify_certification_item(item)
        merged[key] = Assignment(
            raw_name=item,
            display_name=display,
            rubric_kind=rubric_kind,
            sources={"certification"},
        )

    return list(merged.values())


def replace_text_across_runs(paragraph: Paragraph, old: str, new: str) -> bool:
    if not old:
        return False
    if old == new:
        return old in "".join(run.text for run in paragraph.runs)

    replaced = False
    while True:
        full_text = "".join(run.text for run in paragraph.runs)
        start = full_text.find(old)
        if start < 0:
            break

        end = start + len(old)
        positions: list[tuple[int, int]] = []
        cursor = 0
        for run in paragraph.runs:
            next_cursor = cursor + len(run.text)
            positions.append((cursor, next_cursor))
            cursor = next_cursor

        start_run = None
        end_run = None
        start_offset = 0
        end_offset = 0

        for run_index, (run_start, run_end) in enumerate(positions):
            if start_run is None and start < run_end:
                start_run = run_index
                start_offset = start - run_start
            if end <= run_end:
                end_run = run_index
                end_offset = end - run_start
                break

        if start_run is None or end_run is None:
            return replaced

        runs = paragraph.runs
        if start_run == end_run:
            original = runs[start_run].text
            runs[start_run].text = original[:start_offset] + new + original[end_offset:]
        else:
            start_original = runs[start_run].text
            end_original = runs[end_run].text
            runs[start_run].text = start_original[:start_offset] + new
            for run_index in range(start_run + 1, end_run):
                runs[run_index].text = ""
            runs[end_run].text = end_original[end_offset:]

        replaced = True

    return replaced


def set_header(document: Document, code: str) -> None:
    packet_date = academic_date_from_code(code)
    replacements = (
        ("WK", code[0:2]),
        ("DD", code[2:4]),
        ("MM", code[4:6]),
        (
            "Month Date, Year",
            f"{calendar.month_name[packet_date.month]} {packet_date.day}, {packet_date.year}",
        ),
    )

    header_paragraph = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if "Documentation Packet" in paragraph.text and "WK" in paragraph.text
        ),
        None,
    )
    if header_paragraph is None:
        raise RuntimeError("Could not locate the Documentation Packet header.")

    for old, new in replacements:
        if not replace_text_across_runs(header_paragraph, old, new):
            raise RuntimeError(f"Could not find header placeholder {old!r}.")


def set_cell_text(cell: Any, text: str) -> None:
    cell.text = text


def fill_schedule_table(
    document: Document,
    quarter: Any,
    week: Any,
    due_date: Any,
    days: Any,
    workbook_epoch: datetime,
) -> None:
    table = next(
        (
            candidate
            for candidate in document.tables
            if candidate.rows
            and [cell.text.strip() for cell in candidate.rows[0].cells[:4]]
            == ["Quarter", "Week", "Due", "Days"]
        ),
        None,
    )
    if table is None or len(table.rows) < 2:
        raise RuntimeError("Could not locate the schedule table in the template.")

    values = (
        plain_number(quarter),
        plain_number(week),
        format_due_date(due_date, workbook_epoch),
        plain_number(days),
    )
    for cell, value in zip(table.rows[1].cells[:4], values):
        set_cell_text(cell, value)


def find_body_paragraph(document: Document, exact_text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().casefold() == exact_text.casefold():
            return paragraph
    raise RuntimeError(f"Could not locate the {exact_text!r} paragraph.")


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._p.remove(child)


def write_bullet(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    paragraph.add_run("• ")
    paragraph.add_run(text)


def fill_optional_bullet_section(
    document: Document,
    header_text: str,
    next_header_text: str,
    items: list[str],
) -> None:
    """
    Fill a bulleted section between two headers.

    When items is empty, remove the section header and any existing content
    between it and the next header.
    """
    header = find_body_paragraph(document, header_text)
    next_header = find_body_paragraph(document, next_header_text)

    # Remove any old content between the two section headers.
    node = header._p.getnext()
    while node is not None and node is not next_header._p:
        following = node.getnext()
        node.getparent().remove(node)
        node = following

    if not items:
        header._p.getparent().remove(header._p)
        return

    current = header
    for item in items:
        current = insert_paragraph_after(current)
        write_bullet(current, item)


def find_assignments_table(document: Document) -> Table:
    heading = find_body_paragraph(document, "Assignments:")
    node = heading._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return Table(node, heading._parent)
        if node.tag == qn("w:p") and Paragraph(node, heading._parent).text.strip():
            break
        node = node.getnext()
    raise RuntimeError("Could not locate the Assignments table.")


def fill_assignments_table(document: Document, assignments: list[Assignment]) -> None:
    table = find_assignments_table(document)
    if len(table.rows) < 2 or len(table.columns) < 3:
        raise RuntimeError("The Assignments table must have a header and template row.")

    template_row = deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    rows = [("DocPac and Reflections", "Everyone", "")]
    rows.extend((assignment.raw_name, assignment.who, "") for assignment in assignments)

    for assignment_name, who, assigned in rows:
        table._tbl.append(deepcopy(template_row))
        new_row = table.rows[-1]
        set_cell_text(new_row.cells[0], assignment_name)
        set_cell_text(new_row.cells[1], who)
        set_cell_text(new_row.cells[2], assigned)


def rubric_kind_for_heading(text: str) -> str | None:
    normalized = text.strip().casefold()
    if normalized.startswith("nodejs backend:"):
        return "nodejs"
    if normalized.startswith("web browser frontend:"):
        return "web"
    if normalized.startswith("collaboration and professionalism:"):
        return "professional"
    if normalized == "company projects":
        return "company"
    if normalized.startswith("senior assignment:"):
        return "senior"
    if normalized == "performance review":
        return "performance"
    if normalized.startswith("certification:"):
        return "certification"
    return None


def set_table_rows_cant_split(table: Table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def rebuild_rubrics(document: Document, assignments: list[Assignment]) -> None:
    body = document.element.body
    children = list(body.iterchildren())
    prototypes: dict[str, tuple[Any, Any, str]] = {}
    to_remove: list[Any] = []

    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue

        paragraph = Paragraph(child, document)
        heading_text = paragraph.text.strip()
        kind = rubric_kind_for_heading(heading_text)
        if kind is None:
            continue

        table_element = None
        for following in children[index + 1 :]:
            if following.tag == qn("w:tbl"):
                table_element = following
                break
            if following.tag == qn("w:p"):
                following_text = Paragraph(following, document).text.strip()
                if following_text:
                    break

        if table_element is None:
            raise RuntimeError(f"Rubric heading {heading_text!r} is not followed by a table.")

        # Keep only the first prototype of each kind. Senior assignments
        # intentionally use the Collaboration and Professionalism table.
        if kind not in prototypes:
            prototypes[kind] = (
                deepcopy(child),
                deepcopy(table_element),
                heading_text,
            )

        to_remove.extend((child, table_element))

    required_prototypes = {
        "nodejs",
        "web",
        "professional",
        "company",
        "performance",
        "certification",
    }
    missing = sorted(required_prototypes - prototypes.keys())
    if missing:
        raise RuntimeError(
            "Template is missing rubric prototype(s): " + ", ".join(missing)
        )

    for element in dict.fromkeys(to_remove):
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    anchor = find_body_paragraph(document, "DocPac and Reflection")._p

    for assignment in assignments:
        prototype_kind = (
            "professional" if assignment.rubric_kind == "senior" else assignment.rubric_kind
        )
        heading_xml, table_xml, old_heading = prototypes[prototype_kind]
        new_heading_xml = deepcopy(heading_xml)
        new_table_xml = deepcopy(table_xml)

        anchor.addprevious(new_heading_xml)
        anchor.addprevious(new_table_xml)

        new_heading = Paragraph(new_heading_xml, document)
        if not replace_text_across_runs(new_heading, old_heading, assignment.display_name):
            clear_paragraph(new_heading)
            new_heading.add_run(assignment.display_name)
        new_heading.paragraph_format.keep_with_next = True

        new_table = Table(new_table_xml, document)
        set_table_rows_cant_split(new_table)


def get_headers(worksheet: Any) -> dict[str, int]:
    headers: dict[str, int] = {}
    for column, cell in enumerate(worksheet[1], start=1):
        if cell.value is not None:
            headers[str(cell.value).strip().casefold()] = column

    required = (
        "code",
        "quarter",
        "week",
        "date",
        "days",
        "events",
        "juniors",
        "seniors",
        "notes",
    )
    missing = [name for name in required if name not in headers]
    if missing:
        raise RuntimeError(
            "The schedule is missing required column(s): " + ", ".join(missing)
        )

    if "coursework" not in headers:
        raise RuntimeError("The schedule is missing the required Coursework column.")
    return headers


def row_value(row: tuple[Any, ...], headers: dict[str, int], name: str) -> Any:
    return row[headers[name] - 1].value


def coursework_row_value(row: tuple[Any, ...], headers: dict[str, int]) -> Any:
    return row_value(row, headers, "coursework")


def iter_schedule_rows(worksheet: Any, headers: dict[str, int]) -> Iterable[dict[str, Any]]:
    for row_number, row in enumerate(worksheet.iter_rows(min_row=2), start=2):
        code_value = row_value(row, headers, "code")
        if code_value in (None, ""):
            continue

        try:
            code = normalize_code(code_value)
        except ValueError as exc:
            print(f"Skipping row {row_number}: {exc}", file=sys.stderr)
            continue

        junior_items = split_assignment_cell(row_value(row, headers, "juniors"))
        senior_items = split_assignment_cell(row_value(row, headers, "seniors"))
        coursework_items = split_assignment_cell(
            coursework_row_value(row, headers)
        )
        event_items = split_assignment_cell(row_value(row, headers, "events"))
        note_items = split_assignment_cell(row_value(row, headers, "notes"))

        yield {
            "row_number": row_number,
            "code": code,
            "quarter": row_value(row, headers, "quarter"),
            "week": row_value(row, headers, "week"),
            "date": row_value(row, headers, "date"),
            "days": row_value(row, headers, "days"),
            "event_items": event_items,
            "junior_items": junior_items,
            "senior_items": senior_items,
            "coursework_items": coursework_items,
            "note_items": note_items,
            "assignments": build_assignments(
                junior_items,
                senior_items,
                coursework_items,
            ),
        }


def main() -> int:
    args = parse_args()
    root = args.root.resolve()
    schedule_path = args.schedule.resolve()
    template_path = args.template.resolve()

    if not root.is_dir():
        print(f"Root folder does not exist: {root}", file=sys.stderr)
        return 1
    if not schedule_path.is_file():
        print(f"Schedule file does not exist: {schedule_path}", file=sys.stderr)
        return 1
    if not template_path.is_file():
        print(f"Template file does not exist: {template_path}", file=sys.stderr)
        return 1

    workbook = load_workbook(schedule_path, data_only=True)
    if args.sheet:
        if args.sheet not in workbook.sheetnames:
            print(
                f"Worksheet {args.sheet!r} was not found. Available sheets: "
                + ", ".join(workbook.sheetnames),
                file=sys.stderr,
            )
            return 1
        worksheet = workbook[args.sheet]
    else:
        worksheet = workbook.active

    headers = get_headers(worksheet)
    created = 0
    past = 0
    missing_folders = 0

    for item in iter_schedule_rows(worksheet, headers):
        packet_date = academic_date_from_code(item["code"])
        folder = root / f"docpac_{item['code']}"

        if packet_date < args.as_of:
            past += 1
            print(
                f"Skipped {folder.name}: {packet_date.isoformat()} is before "
                f"{args.as_of.isoformat()}."
            )
            continue

        if not folder.is_dir():
            missing_folders += 1
            print(
                f"Skipped {folder.name} (schedule row {item['row_number']}): "
                "folder not found."
            )
            continue

        output_path = folder / f"docpac_{item['code']}.docx"

        if args.dry_run:
            print(
                f"Would rebuild {output_path.relative_to(root)} with "
                f"{len(item['assignments'])} assignment rubric(s) and create "
                f"{len(item['coursework_items'])} Coursework folder(s)."
            )
            continue

        temporary_path = folder / f".~docpac_{item['code']}.tmp.docx"
        temporary_path.unlink(missing_ok=True)
        shutil.copy2(template_path, temporary_path)

        try:
            document = Document(temporary_path)
            set_header(document, item["code"])
            fill_schedule_table(
                document,
                item["quarter"],
                item["week"],
                item["date"],
                item["days"],
                workbook.epoch,
            )
            fill_optional_bullet_section(
                document,
                "Events:",
                "Notes:",
                item["event_items"],
            )
            fill_optional_bullet_section(
                document,
                "Notes:",
                "Assignments:",
                item["note_items"],
            )
            fill_assignments_table(document, item["assignments"])
            rebuild_rubrics(document, item["assignments"])
            document.save(temporary_path)

            output_path.unlink(missing_ok=True)
            temporary_path.replace(output_path)
            create_coursework_folders(folder, item["coursework_items"])
        except Exception:
            temporary_path.unlink(missing_ok=True)
            raise

        print(
            f"Created {output_path.relative_to(root)} "
            f"({len(item['assignments'])} assignments, "
            f"{len(item['coursework_items'])} Coursework folders)."
        )
        created += 1

    print(
        f"\nFinished: {created} rebuilt, {past} past-date skipped, "
        f"{missing_folders} missing-folder skipped."
    )

    if args.dry_run:
        return 0
    return 0 if created or past else 2


if __name__ == "__main__":
    raise SystemExit(main())



